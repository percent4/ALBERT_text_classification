import docx
from docx import Document
import pandas as pd
import numpy as np
import openpyxl
import os


#获取全文文本
def getText(fileName):
    doc = docx.Document(fileName)
    TextList = []
    for paragraph in doc.paragraphs:
        TextList.append(paragraph.text)

    return '\n'.join(TextList)
#删除特定段落
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None


def delWordContent(docx_file='', dest_file=''):
    # 读取文本
    doc = Document(docx_file)
    paragraphs = doc.paragraphs
    i = 0
    flag = False
    for p in paragraphs:
        i += 1
        # print(str(i))
        # print(p.text)
        if p.text.find('需要删除') > -1:
            # print('找到了')
            flag = True
        if flag is True:
            # print('deleting')
            delete_paragraph(p)
    if flag is True:
        # 保存为新文件
        doc.save(dest_file)


# delWordContent(docx_file='H:/temp/test.docx', dest_file='H:/temp/test-new.docx')

def get_titcon(doc, head, context):
    #head = []
    #context = [] 都是集合
    #返回两个集合

    #打印全文
    # for i in range(len(doc.paragraphs)):
    #     print("第" + str(i) + "段的内容是：" + doc.paragraphs[i].text)
    # print("----------------------------------------------------------------")

    # 目标段落，装段落
    #target_para = doc.paragraphs[:-5]
    target_para = doc.paragraphs
    #print(target_para) <docx.text.paragraph.Paragraph object at 0x00000244DB238D30>,
    str = ''
    for para in target_para:
        # print("第" + str(m) + "段的内容是：" + para.text)
        str += para.text

        #print(para.text)
    str = str.replace(u'\u3000', u'').replace('\n', '').replace('\r', '').replace(" ", '')
    #print(str)

    #题目，装题目
    title = doc.paragraphs[0].text


    #print('内容是：' + str + '题目是：' + title)
    #print('内容是：' + str)
    #context = []
    context.append(str)

    #head = []
    head.append(title)
    #df = pd.DataFrame({'title': head, "content": context})

    return head, context

def get_province(path):
    head = []
    context = []
    for root, dirs, files in os.walk(path):
        for file in files:
            path_file = os.path.join(root, file)
            print(path_file)
            doc = Document(path_file)
            head, context = get_titcon(doc, head, context)
    df = pd.DataFrame({'title': head, "content": context})
    return df

if __name__ == '__main__':

    path = r'/\doc2docx\上海\上海市交通委员会关于征集2020年度本市交通运输行业重点节能低碳技术的通知(FBMCLI.14.1658519).docx'

    doc = getText(path)
    print(doc)
# --------------------------------------------------------------------
    doc = Document(path)

    file = r'/\doc2docx\上海'
    # i=1
    # for root, dirs, files in os.walk(file):
    #     for file in files:
    #         path = os.path.join(root, file)
    #         print(path)
    #         doc = Document(path)
    #         print(getText(path))
    #
    #         i +=1
    #         print(i)

    content = []
    for root, dirs, files in os.walk(file):
        for file in files:
            path = os.path.join(root, file)
            print(path)
            doc = Document(path)
            print('---------------------------------------------------------------------------------------------------')
            title = doc.paragraphs[0].text
            print(title)
            txt = ''
            for j in range(len(doc.paragraphs)-4):
                print("第" + str(j) + "段的内容是：" + doc.paragraphs[j].text)

                txt += doc.paragraphs[j].text
            content.append(txt)
            print(txt)


#单独测试有问题省份
    # head=[]
    # context=[]

    # head,context = get_titcon(doc,head,context)
    # print(head)

    # file = r'D:\PycharmProjects\other\doc2docx\上海'
    # i=1
    # for root, dirs, files in os.walk(file):
    #     for file in files:
    #         path = os.path.join(root, file)
    #         print(path)
    #         doc = Document(path)
    #         head, context = get_titcon(doc, head, context)
    #
    #         i +=1
    #         print(i)


    #print(context)

    # df = pd.DataFrame({'title': head, "content": context})
    #
    #
    # print(df)
    #
    # df.to_excel(r'D:\PycharmProjects\other\test_data\广西壮族自治区.xlsx', index=False)



#真
    # path2 = r'D:\PycharmProjects\other\doc2docx'
    #
    # for root, dirs, files in os.walk(path2):
    #     # for file in files:
    #     #     path = os.path.join(root, file)
    #     #     print(path)
    #     #print(root)
    #
    #     for dir in dirs:
    #         file = os.path.join(root,dir)
    #         print(type(file))
    #         print('这个是file：'+ file)
    #
    #         print(type(file))
    #         print('这个是dir：'+dir)
    #
    #         locals()['df_{}'.format(dir)]= get_province(file)
    #         locals()['df_{}'.format(dir)].to_excel(r'D:\PycharmProjects\other\test_data\{}.xlsx'.format(dir), index=False)
    #

            # for x_path in range(file):
            #      print(x_path)



    # str=''
    # title = ''
    # title = doc.paragraphs[0].text
    # for para in doc.paragraphs:
    #
    #     print(para.text)
    #
    #
    #     str += para.text
    #     print(type(str))
    # print("----------------------------------------------------------------")
    # print('内容是：' + str)
    # print("----------------------------------------------------------------")
    # print('题目是：' + title)
    #
    # context = []
    # context.append(str)
    #
    # head = []
    # head.append(title)
    # df = pd.DataFrame({'title': head, "content": context})
    # print(df)
    #
    # df.to_excel(r'D:\PycharmProjects\other\test_data\test.xlsx', index=False)
# -----------------------------------------------------------------------------------------
#最后没用的几段
    # for i in range(len(doc.paragraphs)):
    #     print("第" + str(i) + "段的内容是：" + doc.paragraphs[i].text)
    # print("----------------------------------------------------------------")
    #
    # n = len(doc.paragraphs)
    # print(n)
    # m = n-5
    # print(m)
    # last_para = doc.paragraphs[-7:]
    # print(last_para)
    # str = ''
    # for para in last_para:
    #     #print("第" + str(m) + "段的内容是：" + para.text)
    #     str += para.text
    #
    #     print(para.text)
    #
    # print("----------------------------------------------------------------")
    # print(str)
    #print(str)
# -----------------------------------------------------------------------------------------

    #df.to_excel(r'D:\PycharmProjects\other\test_data\test.xlsx', index=False)
